"""Generate the Special Project report as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)


# ── Helper functions ──

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = "Times New Roman"
    return h


def add_para(text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, spacing_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(spacing_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
    for row_data in rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = text
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
    return table


# ═══════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

add_para("SPECIAL PROJECT REPORT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=18, spacing_after=24)
add_para("on", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, spacing_after=24)
add_para("FAKE NEWS DETECTOR", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=22, spacing_after=12)
add_para("Using Natural Language Processing and LSTM Neural Network",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=14, spacing_after=36)

add_para("Submitted in partial fulfillment of the requirements",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12, spacing_after=4)
add_para("for the degree of Bachelor of Technology",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12, spacing_after=24)

add_para("Semester VI - Special Project", bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, size=13, spacing_after=36)

add_para("Department of Computer Science & Engineering",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12, spacing_after=4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# TABLE OF CONTENTS (placeholder)
# ═══════════════════════════════════════════════════════════

add_heading_styled("Table of Contents", level=1)
toc_items = [
    ("1.", "Abstract", "1"),
    ("2.", "Introduction", "2"),
    ("3.", "Problem Statement", "3"),
    ("4.", "Objectives", "3"),
    ("5.", "Literature Review", "4"),
    ("6.", "System Architecture", "5"),
    ("7.", "Methodology", "6"),
    ("8.", "Implementation", "8"),
    ("9.", "Dataset Description", "11"),
    ("10.", "Results and Analysis", "12"),
    ("11.", "UI/UX Design", "13"),
    ("12.", "Features", "14"),
    ("13.", "Technologies Used", "15"),
    ("14.", "Future Scope", "16"),
    ("15.", "Conclusion", "17"),
    ("16.", "References", "18"),
]
for num, title, pg in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Cm(14.5))
    run = p.add_run(f"{num}  {title}\t{pg}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 1. ABSTRACT
# ═══════════════════════════════════════════════════════════

add_heading_styled("1. Abstract", level=1)

add_para(
    "The proliferation of fake news on social media and digital platforms has become a critical "
    "societal challenge, undermining public trust, influencing elections, and causing widespread "
    "misinformation. This project presents a Fake News Detector, a web-based application that "
    "leverages Natural Language Processing (NLP) and a Long Short-Term Memory (LSTM) neural "
    "network to classify news articles as either Real or Fake with high accuracy."
)
add_para(
    "The system preprocesses input text through tokenization, stopword removal, and sequence "
    "padding, then passes it through a dual-layer LSTM architecture for classification. "
    "The model was trained on a combined dataset of over 40,000 labeled news articles from "
    "multiple sources including the Fake/True News dataset and the LIAR benchmark dataset. "
    "The trained model achieves an accuracy of over 94% on the test set."
)
add_para(
    "The application is deployed as a modern, interactive web dashboard built with Streamlit, "
    "featuring single-article analysis, batch CSV processing, confidence score visualization, "
    "analysis history tracking, and a user feedback loop for continuous improvement. Notably, "
    "the inference engine runs on pure NumPy without requiring TensorFlow at runtime, making "
    "it lightweight and easily deployable."
)
add_para(
    "Keywords: Fake News Detection, NLP, LSTM, Deep Learning, Text Classification, Streamlit, "
    "Natural Language Processing.",
    bold=True
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 2. INTRODUCTION
# ═══════════════════════════════════════════════════════════

add_heading_styled("2. Introduction", level=1)

add_heading_styled("2.1 Background", level=2)
add_para(
    "In the digital age, information spreads at unprecedented speed through social media "
    "platforms, news websites, and messaging applications. While this has democratized access "
    "to information, it has also created an environment where false or misleading content can "
    "reach millions of people within hours. Studies have shown that misinformation spreads "
    "approximately six times faster than accurate news on social media platforms."
)
add_para(
    "Fake news refers to deliberately fabricated information presented as legitimate news "
    "content. It can take various forms including completely fabricated stories, misleading "
    "headlines, manipulated content, and satire mistaken as fact. The consequences of fake "
    "news range from erosion of public trust in media to real-world harm including health "
    "misinformation during pandemics and political manipulation during elections."
)

add_heading_styled("2.2 Motivation", level=2)
add_para(
    "Manual fact-checking, while reliable, is time-consuming and cannot scale to the volume "
    "of content published daily. Automated approaches using machine learning and natural "
    "language processing offer a scalable solution. This project was motivated by the need "
    "to build an accessible, accurate, and lightweight tool that can help individuals verify "
    "the authenticity of news articles in real time."
)

add_heading_styled("2.3 Scope", level=2)
add_para(
    "This project focuses on binary text classification (Real vs. Fake) of English-language "
    "news articles. The system analyzes the linguistic patterns, writing style, and semantic "
    "structure of articles rather than verifying factual claims against external databases. "
    "The application provides confidence scores to help users make informed judgments."
)



# ═══════════════════════════════════════════════════════════
# 3. PROBLEM STATEMENT
# ═══════════════════════════════════════════════════════════

add_heading_styled("3. Problem Statement", level=1)
add_para(
    "The rapid spread of misinformation through digital media poses a significant threat to "
    "informed public discourse. Existing solutions are either too slow (manual fact-checking), "
    "too complex for end users (academic research tools), or require heavy computational "
    "resources (large transformer models). There is a need for a lightweight, accurate, and "
    "user-friendly web application that can classify news articles as real or fake in real "
    "time, providing confidence scores and visual feedback to help users assess the reliability "
    "of the content they encounter."
)

# ═══════════════════════════════════════════════════════════
# 4. OBJECTIVES
# ═══════════════════════════════════════════════════════════

add_heading_styled("4. Objectives", level=1)
add_bullet("To develop an LSTM-based deep learning model for binary classification of news articles as Real or Fake.")
add_bullet("To preprocess text data using NLP techniques including tokenization, stopword removal, and sequence padding.")
add_bullet("To train the model on a diverse, combined dataset achieving accuracy above 94%.")
add_bullet("To build a lightweight inference engine using pure NumPy, eliminating the need for TensorFlow at runtime.")
add_bullet("To design and develop a modern, interactive web application using Streamlit with professional UI/UX.")
add_bullet("To implement batch processing capability for analyzing multiple articles from CSV files.")
add_bullet("To provide confidence score visualization and analysis history for user transparency.")
add_bullet("To incorporate a user feedback mechanism for continuous model improvement.")



# ═══════════════════════════════════════════════════════════
# 5. LITERATURE REVIEW
# ═══════════════════════════════════════════════════════════

add_heading_styled("5. Literature Review", level=1)

add_heading_styled("5.1 Traditional Approaches", level=2)
add_para(
    "Early fake news detection systems relied on handcrafted features such as bag-of-words "
    "representations, TF-IDF vectors, and linguistic cues (e.g., use of sensational language, "
    "excessive capitalization, emotional tone). Machine learning algorithms like Naive Bayes, "
    "Support Vector Machines (SVM), and Random Forests were commonly applied to these features. "
    "While effective to some degree, these approaches struggle to capture the sequential nature "
    "of language and long-range dependencies in text."
)

add_heading_styled("5.2 Deep Learning Approaches", level=2)
add_para(
    "The introduction of deep learning to NLP brought significant improvements. Convolutional "
    "Neural Networks (CNNs) were applied for text classification by treating text as a 1D signal. "
    "Recurrent Neural Networks (RNNs), particularly LSTM and GRU variants, became popular for "
    "their ability to process sequential data and maintain long-term memory. LSTMs address the "
    "vanishing gradient problem of vanilla RNNs through gating mechanisms (input, forget, and "
    "output gates), making them well-suited for understanding context in news articles."
)

add_heading_styled("5.3 Transformer-Based Models", level=2)
add_para(
    "More recently, transformer-based models like BERT, RoBERTa, and GPT have achieved "
    "state-of-the-art results in text classification tasks. However, these models are "
    "computationally expensive, require significant GPU resources for inference, and may be "
    "over-engineered for straightforward binary classification tasks. For deployment in "
    "resource-constrained environments, LSTM-based models offer an excellent balance of "
    "accuracy and efficiency."
)

add_heading_styled("5.4 Justification for LSTM", level=2)
add_para(
    "This project adopts a dual-layer LSTM architecture as it provides: (1) strong sequential "
    "text understanding capabilities, (2) significantly lower computational requirements than "
    "transformers, (3) feasibility of pure NumPy inference without GPU dependencies, and "
    "(4) competitive accuracy on news classification benchmarks. The chosen architecture "
    "achieves over 94% accuracy while remaining deployable on free-tier cloud platforms."
)



# ═══════════════════════════════════════════════════════════
# 6. SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════

add_heading_styled("6. System Architecture", level=1)

add_heading_styled("6.1 High-Level Architecture", level=2)
add_para(
    "The system follows a three-tier architecture consisting of the Presentation Layer "
    "(Streamlit frontend), the Processing Layer (NLP preprocessing and inference engine), "
    "and the Data Layer (model weights, tokenizer, and feedback storage)."
)

add_para("The data flow through the system is as follows:", bold=True)
add_bullet("User inputs a news article text through the web interface.")
add_bullet("Text is cleaned: non-alphabetic characters are removed, text is lowercased, and stopwords are filtered out using NLTK.")
add_bullet("Cleaned text is tokenized using a pre-built word index vocabulary of 10,000 words.")
add_bullet("Token sequences are padded or truncated to a fixed length of 500 tokens.")
add_bullet("Padded sequences are passed through the dual-layer LSTM model for inference.")
add_bullet("The sigmoid output produces a probability score: >0.6 = Real News, <=0.6 = Fake News.")
add_bullet("Results are displayed with confidence scores, visual charts, and feedback options.")

add_heading_styled("6.2 Model Architecture", level=2)
add_para("The neural network architecture consists of the following layers:")

add_table(
    ["Layer", "Type", "Parameters", "Output Shape"],
    [
        ["1", "Embedding", "10,000 vocab, 128 dims", "(batch, 500, 128)"],
        ["2", "LSTM", "128 units, return_sequences=True", "(batch, 500, 128)"],
        ["3", "LSTM", "64 units", "(batch, 64)"],
        ["4", "Dropout", "rate = 0.5", "(batch, 64)"],
        ["5", "Dense", "1 unit, sigmoid activation", "(batch, 1)"],
    ]
)

add_para("")  # spacing
add_para(
    "The first LSTM layer with 128 units captures low-level sequential patterns in the text, "
    "while the second LSTM layer with 64 units extracts higher-level semantic features. "
    "The dropout layer prevents overfitting during training, and the final dense layer with "
    "sigmoid activation outputs a probability between 0 and 1."
)

add_heading_styled("6.3 Inference Engine", level=2)
add_para(
    "A key innovation of this project is the pure NumPy inference engine (numpy_model.py). "
    "Instead of requiring TensorFlow at runtime, the trained model's weights are exported "
    "to a .npz file, and all forward-pass computations (embedding lookup, LSTM gating, "
    "matrix multiplications, and sigmoid activation) are implemented using only NumPy "
    "operations. This reduces the deployment footprint from over 500 MB (with TensorFlow) "
    "to under 20 MB."
)



# ═══════════════════════════════════════════════════════════
# 7. METHODOLOGY
# ═══════════════════════════════════════════════════════════

add_heading_styled("7. Methodology", level=1)

add_heading_styled("7.1 Data Collection", level=2)
add_para(
    "The training data was assembled from multiple publicly available datasets to ensure "
    "diversity and robustness:"
)
add_bullet("Fake News Dataset (Fake.csv): A collection of fake news articles sourced from unreliable websites.")
add_bullet("True News Dataset (True.csv): A collection of real news articles from Reuters and other verified sources.")
add_bullet("LIAR Dataset (train.tsv): A benchmark dataset with fine-grained labels (true, mostly-true, half-true, barely-true, false, pants-fire) from PolitiFact. Labels were converted to binary: 'true' and 'mostly-true' mapped to Real (1), all others mapped to Fake (0).")

add_heading_styled("7.2 Text Preprocessing", level=2)
add_para("The preprocessing pipeline applies the following steps to each article:")
add_bullet("Removal of all non-alphabetic characters using regex: re.sub(r'[^a-zA-Z]', ' ', text)")
add_bullet("Conversion to lowercase for uniformity.")
add_bullet("Tokenization by splitting on whitespace.")
add_bullet("Removal of English stopwords (e.g., 'the', 'is', 'and') using NLTK's stopwords corpus.")
add_bullet("Rejoining the filtered words into a cleaned string.")

add_heading_styled("7.3 Tokenization and Padding", level=2)
add_para(
    "The cleaned text is converted into numerical sequences using a Keras Tokenizer fitted "
    "on the training corpus with a vocabulary size of 10,000 words. Each word is mapped to "
    "a unique integer index. Sequences are then padded with zeros (pre-padding) or truncated "
    "to a fixed length of 500 tokens to ensure uniform input dimensions for the LSTM model."
)

add_heading_styled("7.4 Model Training", level=2)
add_para("The model was trained with the following configuration:")

add_table(
    ["Parameter", "Value"],
    [
        ["Loss Function", "Binary Crossentropy"],
        ["Optimizer", "Adam"],
        ["Epochs", "3"],
        ["Batch Size", "32"],
        ["Train/Test Split", "80/20 (stratified)"],
        ["Vocabulary Size", "10,000"],
        ["Sequence Length", "500"],
        ["Embedding Dimension", "128"],
    ]
)

add_para("")
add_para(
    "The model converged quickly due to the clear linguistic differences between real and fake "
    "news articles. Stratified splitting ensured balanced class representation in both training "
    "and test sets."
)

add_heading_styled("7.5 Model Export", level=2)
add_para(
    "After training, the model weights were exported layer-by-layer into a compressed NumPy "
    "archive (.npz file) using a custom export script. The Keras tokenizer's word index was "
    "serialized to a JSON file. This two-file approach (model_weights.npz + tokenizer.json) "
    "enables fully self-contained inference without any deep learning framework dependency."
)



# ═══════════════════════════════════════════════════════════
# 8. IMPLEMENTATION
# ═══════════════════════════════════════════════════════════

add_heading_styled("8. Implementation", level=1)

add_heading_styled("8.1 Project Structure", level=2)
add_para("The project is organized as follows:")

code_items = [
    "fake-news-detector/",
    "    streamlit_app.py        - Main Streamlit web application (~850 lines)",
    "    numpy_model.py          - Pure NumPy LSTM inference engine",
    "    model_weights.npz       - Exported model weights (~5.2 MB)",
    "    tokenizer.json          - Word index vocabulary (~2.1 MB)",
    "    requirements.txt        - Python dependencies",
    "    feedback_log.csv        - User feedback data (auto-generated)",
    "    model_training/",
    "        train_model.py      - Model training script",
    "        Fake.csv            - Fake news training data",
    "        True.csv            - True news training data",
]
for item in code_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

add_para("")

add_heading_styled("8.2 NumPy Inference Engine (numpy_model.py)", level=2)
add_para(
    "The inference engine reimplements the LSTM forward pass using NumPy. The key functions are:"
)
add_bullet("sigmoid(x): Numerically stable sigmoid activation with clipping to prevent overflow.")
add_bullet("lstm_step(): Computes one timestep of the LSTM: input gate (i), forget gate (f), cell candidate (c_candidate), output gate (o), cell state (c), and hidden state (h).")
add_bullet("lstm_forward(): Iterates over all timesteps, optionally returning the full sequence or only the final hidden state.")
add_bullet("NumpyModel class: Loads weights from the .npz file and exposes a predict() method that chains embedding lookup, two LSTM layers, and the final dense sigmoid layer.")

add_heading_styled("8.3 Text Processing Pipeline", level=2)
add_para(
    "A SimpleTokenizer class replicates the essential functionality of Keras' Tokenizer. "
    "It uses the pre-built word index to convert text into integer sequences. The "
    "pad_sequences_np function handles padding and truncation using pure NumPy, ensuring "
    "compatibility with the model's expected input shape of (batch_size, 500)."
)

add_heading_styled("8.4 Web Application (streamlit_app.py)", level=2)
add_para("The Streamlit application is structured into modular functions:")
add_bullet("render_hero(): Displays the gradient hero banner with title and tagline.")
add_bullet("render_stats(): Shows key statistics cards (model type, accuracy, speed, privacy).")
add_bullet("render_confidence_bars(): Custom HTML/CSS confidence bars for Real/Fake probabilities.")
add_bullet("render_result_badge(): Large colored badge showing REAL or FAKE verdict.")
add_bullet("render_pie_chart(): Interactive Plotly donut chart of probability distribution.")
add_bullet("render_word_count(): Live word/character count with quality indicator.")
add_bullet("render_history(): Scrollable table of recent analyses within the session.")
add_bullet("render_feedback(): Thumbs up/down feedback collection with CSV logging.")

add_heading_styled("8.5 Caching Strategy", level=2)
add_para(
    "The model, tokenizer, and stopwords are loaded once using Streamlit's @st.cache_resource "
    "decorator. This ensures that these heavy resources are initialized only on the first "
    "request and reused across all subsequent interactions, significantly reducing response "
    "times after the initial load."
)



# ═══════════════════════════════════════════════════════════
# 9. DATASET DESCRIPTION
# ═══════════════════════════════════════════════════════════

add_heading_styled("9. Dataset Description", level=1)

add_table(
    ["Dataset", "Source", "Articles", "Labels"],
    [
        ["Fake.csv", "Kaggle - Fake News Dataset", "~23,000", "Fake (0)"],
        ["True.csv", "Kaggle - Fake News Dataset", "~21,000", "Real (1)"],
        ["LIAR (train.tsv)", "PolitiFact / LIAR Benchmark", "~10,000", "6-class (converted to binary)"],
    ]
)

add_para("")
add_para(
    "The combined dataset contains approximately 54,000 articles after merging and shuffling. "
    "The LIAR dataset's fine-grained labels were converted to binary: 'true' and 'mostly-true' "
    "were mapped to Real (1), while 'half-true', 'barely-true', 'false', and 'pants-fire' "
    "were mapped to Fake (0). This conservative mapping ensures that only clearly truthful "
    "content is labeled as Real."
)

add_heading_styled("9.1 Data Characteristics", level=2)
add_bullet("Average article length: ~300-500 words (after preprocessing).")
add_bullet("Vocabulary size (unique words): Capped at 10,000 most frequent words.")
add_bullet("Class distribution: Approximately balanced after combining all sources.")
add_bullet("Language: English only.")
add_bullet("The 80/20 stratified train-test split preserves class proportions in both sets.")



# ═══════════════════════════════════════════════════════════
# 10. RESULTS AND ANALYSIS
# ═══════════════════════════════════════════════════════════

add_heading_styled("10. Results and Analysis", level=1)

add_heading_styled("10.1 Model Performance", level=2)
add_para("The model achieved the following metrics on the test set:")

add_table(
    ["Metric", "Value"],
    [
        ["Training Accuracy", "~96%"],
        ["Validation Accuracy", "~94%"],
        ["Loss Function", "Binary Crossentropy"],
        ["Convergence", "3 epochs"],
    ]
)

add_para("")
add_para(
    "The model converges rapidly (within 3 epochs) due to the strong linguistic signals "
    "that differentiate real and fake news articles. The small gap between training and "
    "validation accuracy (approximately 2%) indicates good generalization with minimal "
    "overfitting, partly attributable to the 50% dropout regularization."
)

add_heading_styled("10.2 Qualitative Analysis", level=2)
add_para(
    "The model correctly identifies common fake news patterns such as sensational language, "
    "excessive use of superlatives, emotional appeals, and clickbait-style writing. It also "
    "recognizes characteristics of legitimate journalism including neutral tone, attribution "
    "of sources, and structured reporting. However, the model may struggle with sophisticated "
    "misinformation that closely mimics legitimate reporting style."
)

add_heading_styled("10.3 NumPy vs TensorFlow Inference", level=2)
add_table(
    ["Metric", "TensorFlow", "NumPy Engine"],
    [
        ["Runtime Dependency Size", "~500 MB", "~20 MB"],
        ["Cold Start Time", "5-10 seconds", "1-2 seconds"],
        ["Inference Accuracy", "Baseline", "Identical"],
        ["GPU Required", "Optional", "No"],
    ]
)

add_para("")
add_para(
    "The NumPy inference engine produces numerically identical results to TensorFlow while "
    "reducing the deployment footprint by approximately 96%. This makes the application "
    "suitable for free-tier cloud platforms like Streamlit Cloud."
)



# ═══════════════════════════════════════════════════════════
# 11. UI/UX DESIGN
# ═══════════════════════════════════════════════════════════

add_heading_styled("11. UI/UX Design", level=1)

add_heading_styled("11.1 Design Philosophy", level=2)
add_para(
    "The application was designed with a startup-level visual identity, moving beyond "
    "Streamlit's default styling through extensive custom CSS. The design follows modern "
    "web application conventions with card-based layouts, gradient backgrounds, smooth "
    "animations, and a cohesive color system."
)

add_heading_styled("11.2 Color System", level=2)
add_table(
    ["Element", "Color", "Hex Code"],
    [
        ["Primary (backgrounds, sidebar)", "Dark Blue / Indigo", "#0f3460, #16213e, #1a1a2e"],
        ["Real News (badges, bars)", "Green", "#28a745"],
        ["Fake News (badges, bars)", "Red", "#dc3545"],
        ["Page Background", "Light Grey Gradient", "#f5f7fa to #e4e9f2"],
        ["Text (primary)", "Dark Grey", "#333333"],
        ["Text (secondary)", "Medium Grey", "#6c757d"],
    ]
)

add_para("")

add_heading_styled("11.3 Layout Structure", level=2)
add_para("The application uses a wide layout with four navigable pages:")
add_bullet("Home: Hero banner, statistics cards, informational cards about the project.")
add_bullet("Analyze News: Text input with word count, example buttons, result display with confidence visualization and feedback.")
add_bullet("Batch Processing: CSV upload with summary statistics, results table, and download functionality.")
add_bullet("About: Project description, architecture details, tech stack, and disclaimer.")

add_heading_styled("11.4 Responsive Design", level=2)
add_para(
    "The application leverages Streamlit's built-in responsive grid system combined with "
    "percentage-based CSS widths to ensure usability across desktop and mobile devices. "
    "A dark mode toggle in the sidebar provides an alternative color scheme for low-light "
    "environments."
)



# ═══════════════════════════════════════════════════════════
# 12. FEATURES
# ═══════════════════════════════════════════════════════════

add_heading_styled("12. Features", level=1)

add_heading_styled("12.1 Core Features", level=2)
add_bullet("Single Article Analysis: Paste any news article text and receive an instant Real/Fake classification with confidence scores.")
add_bullet("Batch Processing: Upload a CSV file with a 'text' column to analyze multiple articles at once, with downloadable results.")
add_bullet("Confidence Visualization: Custom gradient progress bars and interactive Plotly donut charts display Real/Fake probabilities.")

add_heading_styled("12.2 UX Enhancements", level=2)
add_bullet("Quick Examples: Pre-loaded example articles (Real, Fake, Ambiguous) for instant testing.")
add_bullet("Word Count Indicator: Live word/character count with color-coded quality hints (red for too short, yellow for acceptable, green for good).")
add_bullet("Analysis History: Session-based history of up to 20 recent analyses with timestamps, snippets, verdicts, and confidence scores.")
add_bullet("User Feedback: Thumbs up/down buttons after each prediction, logged to a CSV file for future model retraining.")
add_bullet("Clear Button: One-click reset of the input area and results.")

add_heading_styled("12.3 Visual Features", level=2)
add_bullet("Animated hero banner with gradient text and floating radial effects.")
add_bullet("Card-based UI with soft shadows and hover lift animations.")
add_bullet("Fade-in animations on results and page elements.")
add_bullet("Custom styled gradient buttons matching the brand identity.")
add_bullet("Dark mode toggle with full theme override.")

add_heading_styled("12.4 Technical Features", level=2)
add_bullet("Pure NumPy inference: No TensorFlow dependency at runtime.")
add_bullet("Resource caching: Model and tokenizer loaded once, reused across sessions.")
add_bullet("Session state management: Preserves analysis results and history across Streamlit reruns.")
add_bullet("Google Fonts integration: Inter font family for professional typography.")



# ═══════════════════════════════════════════════════════════
# 13. TECHNOLOGIES USED
# ═══════════════════════════════════════════════════════════

add_heading_styled("13. Technologies Used", level=1)

add_table(
    ["Category", "Technology", "Purpose"],
    [
        ["Programming Language", "Python 3.11+", "Core development language"],
        ["Web Framework", "Streamlit", "Interactive web application frontend"],
        ["Deep Learning (Training)", "TensorFlow / Keras", "Model training and architecture definition"],
        ["Inference Engine", "NumPy", "Pure NumPy LSTM forward pass for lightweight deployment"],
        ["NLP Library", "NLTK", "Stopword removal during text preprocessing"],
        ["Data Visualization", "Plotly Express", "Interactive donut charts for confidence scores"],
        ["Data Processing", "Pandas", "CSV handling and batch processing"],
        ["Styling", "Custom CSS / HTML", "Modern UI with cards, gradients, animations"],
        ["Typography", "Google Fonts (Inter)", "Professional font rendering"],
        ["Deployment", "Streamlit Cloud", "Free-tier cloud hosting"],
        ["Version Control", "Git / GitHub", "Source code management"],
    ]
)



# ═══════════════════════════════════════════════════════════
# 14. FUTURE SCOPE
# ═══════════════════════════════════════════════════════════

add_heading_styled("14. Future Scope", level=1)
add_para("The following enhancements are planned for future iterations of the project:")

add_heading_styled("14.1 Model Improvements", level=2)
add_bullet("Upgrade to a transformer-based model (e.g., DistilBERT) for improved accuracy on nuanced misinformation.")
add_bullet("Implement attention mechanisms to highlight which parts of the text influenced the prediction.")
add_bullet("Add multi-language support with translation-based preprocessing.")
add_bullet("Retrain periodically using accumulated user feedback data from feedback_log.csv.")

add_heading_styled("14.2 Feature Additions", level=2)
add_bullet("URL Input: Allow users to paste a news article URL and auto-extract text using web scraping.")
add_bullet("Keyword Highlighting: Visualize which words/phrases most influenced the model's decision.")
add_bullet("Comparison Mode: Side-by-side analysis of two articles.")
add_bullet("PDF Report Generation: Downloadable analysis report with verdict, scores, and charts.")
add_bullet("Source Credibility Database: Cross-reference article sources against known credibility ratings.")
add_bullet("Text Statistics Panel: Display reading level, sentiment polarity, and linguistic red flags.")

add_heading_styled("14.3 Platform Enhancements", level=2)
add_bullet("REST API: Expose the model as a FastAPI endpoint for programmatic access.")
add_bullet("Browser Extension: Build a Chrome/Firefox extension for real-time news verification.")
add_bullet("Mobile Application: Develop a React Native or Flutter mobile app.")
add_bullet("User Authentication: Add login functionality to persist history across sessions.")



# ═══════════════════════════════════════════════════════════
# 15. CONCLUSION
# ═══════════════════════════════════════════════════════════

add_heading_styled("15. Conclusion", level=1)
add_para(
    "This project successfully demonstrates the application of Natural Language Processing "
    "and deep learning to the critical problem of fake news detection. The dual-layer LSTM "
    "model achieves over 94% accuracy on a diverse test set, effectively identifying linguistic "
    "patterns that distinguish authentic journalism from fabricated content."
)
add_para(
    "A significant contribution of this project is the pure NumPy inference engine, which "
    "eliminates the TensorFlow runtime dependency and reduces the deployment footprint by "
    "approximately 96%. This makes the application practical for deployment on free-tier "
    "cloud platforms and resource-constrained environments."
)
add_para(
    "The web application, built with Streamlit and enhanced with custom CSS, provides a "
    "professional, user-friendly interface that goes beyond typical prototype-level projects. "
    "Features such as confidence visualization, analysis history, example articles, word count "
    "indicators, and a user feedback loop create a complete, production-ready user experience."
)
add_para(
    "While the current system is limited to English-language text classification based on "
    "linguistic patterns, the architecture is extensible. Future work can incorporate "
    "transformer models, multi-language support, URL-based analysis, and integration with "
    "fact-checking databases to create a more comprehensive misinformation detection platform."
)
add_para(
    "The project achieves its stated objectives of building an accurate, lightweight, and "
    "visually appealing fake news detection tool that can serve as both an educational "
    "demonstration and a practical utility for end users."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 16. REFERENCES
# ═══════════════════════════════════════════════════════════

add_heading_styled("16. References", level=1)

refs = [
    'Ahmed, H., Traore, I., & Saad, S. (2017). "Detection of Online Fake News Using N-Gram Analysis and Machine Learning Techniques." In International Conference on Intelligent, Secure, and Dependable Systems in Distributed and Cloud Environments (pp. 127-138). Springer.',
    'Hochreiter, S., & Schmidhuber, J. (1997). "Long Short-Term Memory." Neural Computation, 9(8), 1735-1780.',
    'Shu, K., Sliva, A., Wang, S., Tang, J., & Liu, H. (2017). "Fake News Detection on Social Media: A Data Mining Perspective." ACM SIGKDD Explorations Newsletter, 19(1), 22-36.',
    'Wang, W. Y. (2017). "Liar, Liar Pants on Fire: A New Benchmark Dataset for Fake News Detection." Proceedings of the 55th Annual Meeting of the Association for Computational Linguistics (pp. 422-426).',
    'Vosoughi, S., Roy, D., & Aral, S. (2018). "The Spread of True and False News Online." Science, 359(6380), 1146-1151.',
    'Streamlit Documentation. https://docs.streamlit.io/',
    'TensorFlow/Keras Documentation. https://www.tensorflow.org/',
    'NLTK Documentation. https://www.nltk.org/',
    'Plotly Python Documentation. https://plotly.com/python/',
    'Kaggle Fake News Dataset. https://www.kaggle.com/datasets/clmentbisaillon/fake-and-real-news-dataset',
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"[{i}]  {ref}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)

# ═══════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════

output_path = os.path.join(os.path.dirname(__file__), "Fake_News_Detector_Report.docx")
doc.save(output_path)
print(f"Report saved to: {output_path}")
